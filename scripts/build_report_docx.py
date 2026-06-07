"""Build a DOCX report from report/experiment_report.md."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_FILE = ROOT / "report" / "experiment_report.md"
OUTPUT_DIR = ROOT / "output" / "doc"
OUTPUT_FILE = OUTPUT_DIR / "text_mining_report.docx"

IMAGE_RE = re.compile(r"!\[(?P<alt>.*?)\]\((?P<path>.*?)\)")


def set_run_font(run, name: str = "宋体", size: int = 11, bold: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def normalize_inline(text: str) -> str:
    return text.replace("`", "").strip()


def apply_doc_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)


def add_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, name="黑体", size=16, bold=True)


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {min(level, 3)}"
    run = paragraph.add_run(text)
    size = 14 if level == 1 else 12
    set_run_font(run, name="黑体", size=size, bold=True)


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(22)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(normalize_inline(text))
    set_run_font(run, size=11)


def add_list_item(doc: Document, text: str, ordered: bool = False) -> None:
    style = "List Number" if ordered else "List Bullet"
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(normalize_inline(text))
    set_run_font(run, size=11)


def add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Pt(18)
        run = paragraph.add_run(line)
        set_run_font(run, name="Courier New", size=9)


def split_table_row(line: str) -> list[str]:
    cells = line.strip().strip("|").split("|")
    return [normalize_inline(cell) for cell in cells]


def is_separator_row(line: str) -> bool:
    content = line.strip().strip("|")
    parts = [part.strip() for part in content.split("|")]
    return bool(parts) and all(set(part) <= {"-", ":"} and "-" in part for part in parts)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=1, cols=width)
    table.style = "Table Grid"

    for col_idx in range(width):
        text = rows[0][col_idx] if col_idx < len(rows[0]) else ""
        run = table.rows[0].cells[col_idx].paragraphs[0].add_run(text)
        set_run_font(run, name="黑体", size=9, bold=True)

    for row in rows[1:]:
        cells = table.add_row().cells
        for col_idx in range(width):
            text = row[col_idx] if col_idx < len(row) else ""
            run = cells[col_idx].paragraphs[0].add_run(text)
            set_run_font(run, size=9)


def resolve_image_path(markdown_path: str) -> Path:
    raw = markdown_path.strip()
    candidate = Path(raw)
    if candidate.is_absolute():
        return candidate
    return (SOURCE_FILE.parent / candidate).resolve()


def add_image(doc: Document, alt: str, markdown_path: str) -> None:
    image_path = resolve_image_path(markdown_path)
    if not image_path.exists():
        add_body_paragraph(doc, f"{alt}：图片文件未找到，路径 {markdown_path}")
        return
    doc.add_picture(str(image_path), width=Inches(6.0))
    if alt:
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run(alt)
        set_run_font(run, size=10)


def build_docx() -> Path:
    lines = SOURCE_FILE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    apply_doc_style(doc)

    paragraph_buffer: list[str] = []
    in_code = False
    code_buffer: list[str] = []
    idx = 0

    def flush_paragraph() -> None:
        if paragraph_buffer:
            add_body_paragraph(doc, " ".join(part.strip() for part in paragraph_buffer))
            paragraph_buffer.clear()

    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_buffer)
                code_buffer.clear()
                in_code = False
            else:
                flush_paragraph()
                in_code = True
            idx += 1
            continue

        if in_code:
            code_buffer.append(line)
            idx += 1
            continue

        if not stripped:
            flush_paragraph()
            idx += 1
            continue

        image_match = IMAGE_RE.fullmatch(stripped)
        if image_match:
            flush_paragraph()
            add_image(doc, image_match.group("alt"), image_match.group("path"))
            idx += 1
            continue

        if stripped.startswith("#"):
            flush_paragraph()
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            if level == 1:
                add_title(doc, text)
            else:
                add_heading(doc, text, level - 1)
            idx += 1
            continue

        if stripped.startswith("|") and idx + 1 < len(lines) and is_separator_row(lines[idx + 1]):
            flush_paragraph()
            table_rows = [split_table_row(stripped)]
            idx += 2
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                table_rows.append(split_table_row(lines[idx]))
                idx += 1
            add_table(doc, table_rows)
            continue

        ordered_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if ordered_match:
            flush_paragraph()
            add_list_item(doc, ordered_match.group(1), ordered=True)
            idx += 1
            continue

        if stripped.startswith("- "):
            flush_paragraph()
            add_list_item(doc, stripped[2:], ordered=False)
            idx += 1
            continue

        paragraph_buffer.append(stripped)
        idx += 1

    flush_paragraph()
    if code_buffer:
        add_code_block(doc, code_buffer)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_FILE)
    return OUTPUT_FILE


if __name__ == "__main__":
    print(build_docx())
